#!/usr/bin/env python3
"""Generate Dr. Sam Quill CV as a Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup: A4, 2.5cm margins ────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width    = Cm(21)
sec.page_height   = Cm(29.7)
sec.top_margin    = Cm(2.2)
sec.bottom_margin = Cm(2.2)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)

GOLD = RGBColor(0x8b, 0x6f, 0x47)
DARK = RGBColor(0x3a, 0x2a, 0x18)
MID  = RGBColor(0x55, 0x42, 0x2e)

# ── Helpers ───────────────────────────────────────────────────────────────────
def sp(para, before=0, after=4):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)

def add_rule(doc, color='C8A96E', before=4, after=4):
    p = doc.add_paragraph()
    sp(p, before=before, after=after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def section_heading(doc, text):
    """Section title in small-caps style with gold rule underneath."""
    p = doc.add_paragraph()
    sp(p, before=10, after=0)
    run = p.add_run(text.upper())
    run.font.name  = 'Garamond'
    run.font.size  = Pt(10.5)
    run.font.bold  = True
    run.font.color.rgb = GOLD
    # Thin gold rule immediately below heading
    add_rule(doc, color='D4B483', before=0, after=3)
    return p

def body(doc, text, italic=False, bold=False, before=0, after=4):
    p = doc.add_paragraph()
    sp(p, before=before, after=after)
    r = p.add_run(text)
    r.font.name   = 'Garamond'
    r.font.size   = Pt(11)
    r.font.italic = italic
    r.font.bold   = bold
    return p, r

def rich_para(doc, parts, before=0, after=5, indent=None):
    """parts = list of (text, italic, bold)"""
    p = doc.add_paragraph()
    sp(p, before=before, after=after)
    if indent:
        p.paragraph_format.left_indent = indent
    for text, italic, bold in parts:
        r = p.add_run(text)
        r.font.name   = 'Garamond'
        r.font.size   = Pt(11)
        r.font.italic = italic
        r.font.bold   = bold
    return p

def bullet(doc, parts_or_text, level=0, after=3):
    """Bullet paragraph with hanging indent. parts_or_text may be a string or list of (text,italic,bold)."""
    p = doc.add_paragraph()
    sp(p, before=0, after=after)
    left = Inches(0.22 + level * 0.22)
    hang = Inches(0.18)
    p.paragraph_format.left_indent    = left
    p.paragraph_format.first_line_indent = -hang
    # bullet character run
    r_bul = p.add_run('– ')   # en-dash + en-space
    r_bul.font.name  = 'Garamond'
    r_bul.font.size  = Pt(11)
    r_bul.font.color.rgb = GOLD
    if isinstance(parts_or_text, str):
        r = p.add_run(parts_or_text)
        r.font.name = 'Garamond'; r.font.size = Pt(11)
    else:
        for text, italic, bold in parts_or_text:
            r = p.add_run(text)
            r.font.name = 'Garamond'; r.font.size = Pt(11)
            r.font.italic = italic; r.font.bold = bold
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p, before=0, after=2)
r = p.add_run('DR. SAM QUILL')
r.font.name = 'Garamond'; r.font.size = Pt(24); r.font.bold = True
r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p, before=0, after=2)
r = p.add_run('Private Tutor in English, History, Philosophy & Critical Thinking')
r.font.name = 'Garamond'; r.font.size = Pt(12); r.font.italic = True
r.font.color.rgb = MID

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p, before=0, after=4)
r = p.add_run('sam@samquill.com  ·  www.samquill.com')
r.font.name = 'Garamond'; r.font.size = Pt(10)
r.font.color.rgb = GOLD

add_rule(doc, before=0, after=6)

# ── Intro ─────────────────────────────────────────────────────────────────────
body(doc,
    'I have worked for many years as a private tutor in English, History, Philosophy and '
    'Critical Thinking, as well as academic and creative writing. Since September 2019 it '
    'has been my full-time occupation.',
    before=0, after=10)

# ═══════════════════════════════════════════════════════════════════════════════
# SPECIALISM
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Specialism: Public School Scholarship Exams (13+)')

body(doc,
    'I specialise in tutoring students who wish to sit for a scholarship at one of the major '
    'British public schools; most often Winchester College, Eton College and Westminster School.')

body(doc, 'Since 2019, many of my students have won scholarships:', after=4)

for school, count in [('Westminster School',7),('Winchester College',5),('Eton College',2),('Harrow School',1)]:
    p = doc.add_paragraph()
    sp(p, before=0, after=2)
    p.paragraph_format.left_indent = Inches(0.3)
    rn = p.add_run(f'{count} ')
    rn.font.name = 'Garamond'; rn.font.size = Pt(13); rn.font.bold = True
    rn.font.color.rgb = GOLD
    rs = p.add_run(school)
    rs.font.name = 'Garamond'; rs.font.size = Pt(11)

body(doc,
    'Many other students have won non-scholarship places, including alongside those schools '
    'named above at St. Paul’s Girls’ School.',
    before=4)

# ═══════════════════════════════════════════════════════════════════════════════
# ACADEMIC QUALIFICATIONS
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Academic Qualifications')

quals = [
    ('Ph.D.',        'Queen Mary, University of London, 2014–18',
     ['‘Fast Influencings: Figuring Necessity in Percy Bysshe Shelley’',
      'Supervisor: Prof. Paul Hamilton  ·  Secondary supervisor: Dr. James Vigus',
      'Funding: Principal’s Studentship']),
    ('M.Phil.',      'University of Cambridge (Gonville and Caius College), 2013–14',
     ['Eighteenth Century and Romantic Studies: Merit (73)']),
    ('B.A. (Hons.)', 'University of Oxford (Christ Church), 2009–12',
     ['English Language and Literature: First Class']),
]

for degree, institution, details in quals:
    p = doc.add_paragraph()
    sp(p, before=4, after=1)
    r1 = p.add_run(degree + ' ')
    r1.font.name = 'Garamond'; r1.font.size = Pt(11); r1.font.bold = True
    r2 = p.add_run(institution)
    r2.font.name = 'Garamond'; r2.font.size = Pt(11)
    for d in details:
        p2 = doc.add_paragraph()
        sp(p2, before=0, after=1)
        p2.paragraph_format.left_indent = Inches(0.25)
        r = p2.add_run(d)
        r.font.name = 'Garamond'; r.font.size = Pt(10.5); r.font.italic = True

# ═══════════════════════════════════════════════════════════════════════════════
# DISCRETION
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Discretion (High-Profile Clients)')

body(doc,
    'I have experience working with public figures and their families, for whom privacy is '
    'paramount; discretion is part of what I offer. I have excellent rates of success under '
    'such circumstances. References can be provided, in confidence, on request.')

# ═══════════════════════════════════════════════════════════════════════════════
# UNIVERSITY ENTRANCE
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'University Entrance')

for item in [
    'One of my students successfully won a place to read English Language and Literature at Keble College, Oxford.',
    'Another won a place to read Philosophy, Politics and Economics at Balliol College, Oxford.',
    'A third won a place to read English Language and Literature at Christ Church, Oxford.',
    'I can help students craft successful personal statements for entrance to U.K. universities, with particular experience in Oxbridge applications.',
    'I have experience in assisting applications for postgraduate study: writing cover-letters, personal statements and research proposals.',
]:
    bullet(doc, item)

# ═══════════════════════════════════════════════════════════════════════════════
# A-LEVEL / GCSE / KS2
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'A-Level')
bullet(doc, 'The AQA syllabus in A-Level Philosophy.')

section_heading(doc, 'GCSE')
for item in [
    'The AQA syllabus in Religious Studies.',
    'The Edexcel International GCSE in English Literature.',
    'Shakespeare for GCSE English.',
    'General tuition in History.',
]:
    bullet(doc, item)

section_heading(doc, 'Key Stage Two')
for item in [
    'English poetry and prose, including for internal school exams.',
    'Religious Studies, including for internal school exams.',
    'History, including for internal school exams.',
    'General Studies, including for internal school exams.',
    'Creative writing, including for internal school exams.',
]:
    bullet(doc, item)

# ═══════════════════════════════════════════════════════════════════════════════
# ESSAY PRIZES
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Essay Prizes')
body(doc,
    'I have, for the past three years, both individually and in classes, helped students to '
    'prepare their submissions for the John Locke Institute’s Global Essay Prize.')
body(doc,
    'Preparation involves a combination of teaching around the broader subject of their question '
    '(philosophical or political), recommending reading, and helping them structure their writing '
    'for the final submission.')

# ═══════════════════════════════════════════════════════════════════════════════
# CRITICAL THINKING
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Critical Thinking (Interview Preparation, 11+)')
body(doc, 'I am experienced in preparing younger students for their school interviews.')
body(doc,
    'Typically, this involves learning and practising intelligent spoken responses to news '
    'articles, and paintings and photographs.')

# ═══════════════════════════════════════════════════════════════════════════════
# CREATIVE WRITING
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Creative Writing')
body(doc,
    'I am a published poet and librettist with a good deal of experience in teaching students '
    'how to improve their creative writing, both in the specific context of examinations and '
    'their broader literary attempts and endeavours.')

# ═══════════════════════════════════════════════════════════════════════════════
# SCHOOL TEACHING
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'School Teaching')
body(doc,
    'Auxiliar de Conversación at Colegio León Felipe, Madrid (January–June 2019). '
    'I was a language assistant in a primary school just outside of Madrid.')

# ═══════════════════════════════════════════════════════════════════════════════
# TEACHING QUALIFICATION & HE EXPERIENCE
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Teaching Qualification & Higher Education Teaching Experience')

p = doc.add_paragraph()
sp(p, before=2, after=1)
r1 = p.add_run('Certificate in Learning and Teaching (CILT): ')
r1.font.name = 'Garamond'; r1.font.size = Pt(11); r1.font.bold = True
r2 = p.add_run('Distinction (2018)')
r2.font.name = 'Garamond'; r2.font.size = Pt(11)

body(doc, 'I am an Associate Fellow of the Higher Education Academy.')

body(doc,
    'From 2015 to 2019, I was a seminar leader for undergraduate courses in English Literature '
    'at Queen Mary, University of London. Courses taught:',
    after=3)

for title, years, year_group in [
    ('‘Poetry’',                                       '2016, 2017, 2018, 2019',  'First Year'),
    ('‘English in Practice’',                          '2019',                     'First Year'),
    ('‘Romantics and Revolutionaries’',                '2016–17, 2017–18','Second Year'),
    ('‘Representing London: Writing the Eighteenth Century City’', '2017–18', 'Second Year'),
    ('‘Imagination and Knowledge’',                    '2016',                     'Second Year'),
    ('‘Reading, Theory and Interpretation’',           '2015–16',             'First Year'),
]:
    bullet(doc, [(title, True, False), (f',  {years}  ({year_group})', False, False)])

# ═══════════════════════════════════════════════════════════════════════════════
# GENERAL PUBLICATIONS
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'General Publications')

# TLS reviews (parent bullet)
bullet(doc, [('Two reviews published in the ', False, False),
             ('Times Literary Supplement', True, False),
             (':', False, False)], after=2)

for review in [
    'A review of Don Paterson’s poetry collection, ‘The Arctic’ (February, 2023).',
    'A review of Deja Whitehouse’s book on tarot and the occult, ‘The Lady and the Beast’ (June, 2025).',
]:
    bullet(doc, review, level=1)

bullet(doc, [('A chapter on the Peterloo Massacre in ', False, False),
             ('English Radical History', True, False),
             (', ed. Stephen Basdeo (Barnsley: Pen and Sword, 2021), 8,000 words.', False, False)])

bullet(doc, [('A ‘defence of poetry’: ‘How Heavy’s the Albatross?’ in ', False, False),
             ('New Defences of Poetry', True, False),
             (' (Newcastle Centre for the Literary Arts, 2021).', False, False)])

bullet(doc, [('Editorial advisor, overseeing the curation and reissuing of the Percy Bysshe Shelley entry in the ', False, False),
             ('Nineteenth-Century Literature Criticism', True, False),
             (' series (Layman Poupard, 2022).', False, False)])

# ═══════════════════════════════════════════════════════════════════════════════
# POETRY PUBLICATIONS
# ═══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Poetry Publications')

poetry = [
    [('‘Cleopatra’, a submitted and accepted commission for an anthology of poems on '
      '‘Shakespeare’s Women’ (forthcoming with Broken Sleep Books in 2026).', False, False)],
    [('‘She was so unlike him. She was so much like us’, commissioned for ', False, False),
     ('Anne-thology: Poems Re-Presenting Anne Shakespeare', True, False),
     (' (Broken Sleep Books, 2023).', False, False)],
    [('', False, False),
     ('Hey Ho The White Swan By God I Am Thy Man', True, False),
     (' (a pamphlet on Broken Sleep Books, 2022).', False, False)],
    [('‘The Lost Decade’ and ‘The Ship of Theseus’ in ', False, False),
     ('Poetry Birmingham Literary Journal', True, False),
     (' (Autumn/Winter, 2020).', False, False)],
    [('‘Damnatio Memoriae’, poem of the week on the London Review of Books Bookshop Blog (August 2020).', False, False)],
    [('‘Coiner’ in ', False, False),
     ('BLACKBOX MANIFOLD', True, False),
     (' (No.․24, Summer 2020).', False, False)],
    [('Various poems in ', False, False),
     ('The Next Review', True, False),
     (' (Aug–Sep 2016).', False, False)],
    [('‘14 Sonnets’ in ', False, False),
     ('PN Review', True, False),
     (' (May–June, 2016).', False, False)],
]

for parts in poetry:
    # strip empty leading parts
    filtered = [(t, i, b) for t, i, b in parts if t]
    if filtered:
        bullet(doc, filtered)

# ── Closing ───────────────────────────────────────────────────────────────────
add_rule(doc, before=10, after=6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p, before=0, after=2)
r = p.add_run('Please get in touch to discuss needs, availability, and rates.')
r.font.name = 'Garamond'; r.font.size = Pt(11); r.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p, before=0, after=0)
r = p.add_run('sam@samquill.com')
r.font.name = 'Garamond'; r.font.size = Pt(11)
r.font.color.rgb = GOLD

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/samquill/Documents/GitHub/samquill00-rgb.github.io/Sam_Quill_CV.docx'
doc.save(out)
print(f'Saved: {out}')
